import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import fitz  # PyMuPDF
import os
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Inches
import markdown
import tempfile
import shutil

class PDFMDMerger:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF & Markdown 合并工具")
        self.root.geometry("1000x800")
        
        # 存储临时文件和图片路径
        self.temp_dir = tempfile.mkdtemp()
        self.image_paths = []
        self.current_image_index = 0
        
        self.setup_ui()
    
    def setup_ui(self):
        # 创建主框架
        self.main_frame = ttk.Frame(self.root)
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 创建按钮区域
        self.button_frame = ttk.Frame(self.main_frame)
        self.button_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 添加按钮
        self.pdf_button = ttk.Button(self.button_frame, text="选择PDF文件", command=self.load_pdf)
        self.pdf_button.pack(side=tk.LEFT, padx=5)
        
        self.md_button = ttk.Button(self.button_frame, text="导入MD文件", command=self.load_markdown)
        self.md_button.pack(side=tk.LEFT, padx=5)
        
        self.export_button = ttk.Button(self.button_frame, text="导出DOCX", command=self.export_docx)
        self.export_button.pack(side=tk.LEFT, padx=5)
        
        # 创建Notebook（标签页）
        self.notebook = ttk.Notebook(self.main_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True)
        
        # 创建图片显示页
        self.image_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.image_frame, text="PDF预览")
        
        # 创建滚动条和画布
        self.canvas = tk.Canvas(self.image_frame)
        self.scrollbar = ttk.Scrollbar(self.image_frame, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = ttk.Frame(self.canvas)
        
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )
        
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        
        self.scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)
        
        # 创建Markdown编辑页
        self.md_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.md_frame, text="Markdown编辑")
        
        self.md_text = tk.Text(self.md_frame)
        self.md_text.pack(fill=tk.BOTH, expand=True)
    
    def load_pdf(self):
        try:
            pdf_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
            if not pdf_path:
                return
                
            # 清理之前的图片
            self.image_paths = []
            for file in os.listdir(self.temp_dir):
                os.remove(os.path.join(self.temp_dir, file))
            
            # 清理之前的图片显示
            for widget in self.scrollable_frame.winfo_children():
                widget.destroy()
            
            # 打开PDF并转换为图片
            pdf_document = fitz.open(pdf_path)
            
            # 设置合适的显示尺寸
            display_width = 800  # 显示宽度
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                pix = page.get_pixmap()
                image_path = os.path.join(self.temp_dir, f"page_{page_num+1}.png")
                pix.save(image_path)
                self.image_paths.append(image_path)
                
                # 创建并显示图片
                image = Image.open(image_path)
                # 计算等比例缩放后的高度
                ratio = display_width / image.width
                new_height = int(image.height * ratio)
                image = image.resize((display_width, new_height), Image.Resampling.LANCZOS)
                
                photo = ImageTk.PhotoImage(image)
                label = ttk.Label(self.scrollable_frame, image=photo)
                label.image = photo  # 保持引用
                label.pack(pady=5)
            
            pdf_document.close()
            messagebox.showinfo("成功", "PDF文件加载完成！")
            
        except Exception as e:
            messagebox.showerror("错误", f"PDF处理出错：{str(e)}")
    
    def show_current_image(self):
        if not self.image_paths:
            return
            
        image_path = self.image_paths[self.current_image_index]
        image = Image.open(image_path)
        
        # 调整图片大小以适应窗口
        display_size = (800, 600)
        image.thumbnail(display_size, Image.Resampling.LANCZOS)
        
        photo = ImageTk.PhotoImage(image)
        self.image_label.configure(image=photo)
        self.image_label.image = photo
    
    def prev_image(self):
        if self.current_image_index > 0:
            self.current_image_index -= 1
            self.show_current_image()
    
    def next_image(self):
        if self.current_image_index < len(self.image_paths) - 1:
            self.current_image_index += 1
            self.show_current_image()
    
    def load_markdown(self):
        try:
            md_path = filedialog.askopenfilename(filetypes=[("Markdown files", "*.md")])
            if not md_path:
                return
                
            with open(md_path, 'r', encoding='utf-8') as f:
                content = f.read()
                self.md_text.delete('1.0', tk.END)
                self.md_text.insert('1.0', content)
                
        except Exception as e:
            messagebox.showerror("错误", f"Markdown文件加载出错：{str(e)}")
    
    def export_docx(self):
        try:
            if not self.image_paths:
                messagebox.showwarning("警告", "请先加载PDF文件！")
                return
                
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")]
            )
            if not save_path:
                return
            
            # 创建Word文档
            doc = Document()
            
            # 添加Markdown内容（直接添加文本，不转换为HTML）
            md_content = self.md_text.get('1.0', tk.END).strip()
            if md_content:
                doc.add_paragraph(md_content)
            
            # 添加图片，使用正确的 Inches 引用
            for image_path in self.image_paths:
                doc.add_picture(image_path, width=Inches(6))  # 修改这行
            
            # 保存文档
            doc.save(save_path)
            messagebox.showinfo("成功", "文档导出成功！")
            
        except Exception as e:
            messagebox.showerror("错误", f"导出文档时出错：{str(e)}")
    
    def __del__(self):
        # 清理临时文件夹
        try:
            shutil.rmtree(self.temp_dir)
        except:
            pass

if __name__ == "__main__":
    root = tk.Tk()
    app = PDFMDMerger(root)
    root.mainloop()
