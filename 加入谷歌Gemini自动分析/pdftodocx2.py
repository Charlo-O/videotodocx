import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox
import fitz  # PyMuPDF
import requests
from PIL import Image
import io
from datetime import datetime
from docx import Document
import threading
import markdown
from docx.shared import Inches
import base64
from PIL import Image, ImageTk  # 添加 ImageTk
import time
import numpy as np

class PDFAnalyzer:
    def process_md(self):
        """处理选定的 Markdown 文件"""
        md_path = filedialog.askopenfilename(filetypes=[("Markdown Files", "*.md")])
        if not md_path:
            return

        self.md_text_area.delete("1.0", tk.END)
        self.md_text_area.insert(tk.END, f"正在处理 Markdown: {md_path}\n")
        self.md_analysis_results = []

        def process_md_thread():
            try:
                with open(md_path, 'r', encoding='utf-8') as file:
                    md_content = file.read()
                
                # 获取当前PDF分析结果中的页数
                pdf_pages_count = len(self.pdf_analysis_results)
                
                if pdf_pages_count == 0:
                    self.md_text_area.insert(tk.END, "请先处理PDF文件以确定分段数量\n")
                    return
                
                # 首先优化整段文字
                headers = {
                    "Content-Type": "application/json",
                }
                
                optimization_data = {
                    "contents": [{
                        "parts":[{
                            "text": f"""请对以下文本进行优化和整理，使其更加清晰和准确。联系上下文，纠正里面的错别字，符合事实，前后文相互呼应，合理运用标点符号，使其更加通顺。
                            文本内容需要被分成{pdf_pages_count}个段落，每个段落的内容要完整且相关。
                            直接返回优化后的文本内容，不要添加任何额外的解释或标记。

                            原文本：
                            {md_content}"""
                        }]
                    }],
                    "generationConfig": {
                        "temperature": 0.7,
                        "maxOutputTokens": 2000,
                    }
                }

                response = requests.post(
                    f"{GOOGLE_API_URL}?key={GOOGLE_API_KEY}",
                    headers=headers,
                    json=optimization_data,
                    timeout=30
                )
                response.raise_for_status()
                
                if response.status_code == 200:
                    result = response.json()
                    if "candidates" in result and len(result["candidates"]) > 0:
                        optimized_content = result["candidates"][0]["content"]["parts"][0]["text"]
                        
                        # 让模型将优化后的内容分成指定数量的段落
                        segmentation_data = {
                            "contents": [{
                                "parts":[{
                                    "text": f"""请将以下文本准确地分成{pdf_pages_count}段，每段都应该是完整的意思单元。
                                    直接返回分段后的文本，用<segment>标签标记每段内容。
                                    例如：
                                    <segment>第一段内容</segment>
                                    <segment>第二段内容</segment>
                                    
                                    待分段文本：
                                    {optimized_content}"""
                                }]
                            }],
                            "generationConfig": {
                                "temperature": 0.7,
                                "maxOutputTokens": 2000,
                            }
                        }

                        segment_response = requests.post(
                            f"{GOOGLE_API_URL}?key={GOOGLE_API_KEY}",
                            headers=headers,
                            json=segmentation_data,
                            timeout=30
                        )
                        segment_response.raise_for_status()
                        
                        if segment_response.status_code == 200:
                            segment_result = segment_response.json()
                            if "candidates" in segment_result and len(segment_result["candidates"]) > 0:
                                segmented_content = segment_result["candidates"][0]["content"]["parts"][0]["text"]
                                
                                # 解析分段内容
                                import re
                                segments = re.findall(r'<segment>(.*?)</segment>', segmented_content, re.DOTALL)
                                
                                # 处理每个分段
                                for i, segment in enumerate(segments, 1):
                                    segment = segment.strip()
                                    self.md_analysis_results.append((i, segment, None))
                                    
                                    self.md_text_area.insert(tk.END, f"\n第{i}段内容:\n{segment}\n")
                                    self.md_text_area.insert(tk.END, "-" * 50 + "\n")
                                
                                self.md_text_area.insert(tk.END, "\n分段完成！\n")
                            else:
                                self.md_text_area.insert(tk.END, "分段处理返回格式异常\n")
                        else:
                            self.md_text_area.insert(tk.END, f"分段处理失败：HTTP {segment_response.status_code}\n")
                    else:
                        self.md_text_area.insert(tk.END, "优化处理返回格式异常\n")
                else:
                    self.md_text_area.insert(tk.END, f"优化处理失败：HTTP {response.status_code}\n")
                        
            except Exception as e:
                self.md_text_area.insert(tk.END, f"\n处理失败: {str(e)}\n")
                import traceback
                self.md_text_area.insert(tk.END, f"详细错误：{traceback.format_exc()}\n")

        threading.Thread(target=process_md_thread, daemon=True).start()

    def save_md_analysis(self):
        """保存 Markdown 分析结果"""
        if not self.md_analysis_results:
            self.md_text_area.insert(tk.END, "没有可供下载的分析结果\n")
            return

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_filename = f"markdown_analysis_{timestamp}.docx"
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=default_filename,
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return

        try:
            doc = Document()
            doc.add_heading('Markdown文档分析报告', 0)

            for _, result, _ in self.md_analysis_results:
                doc.add_paragraph(result)
                doc.add_paragraph('_' * 50)

            doc.save(file_path)
            self.md_text_area.insert(tk.END, f"\n分析结果已保存至: {file_path}\n")
            
        except Exception as e:
            self.md_text_area.insert(tk.END, f"\n保存文件失败: {str(e)}\n")
            import traceback
            self.md_text_area.insert(tk.END, f"详细错误：{traceback.format_exc()}\n")

    def process_pdf(self):
        """处理选定的 PDF 文件，提取图像并发送给大模型分析"""
        pdf_path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
        if not pdf_path:
            return

        self.pdf_text_area.delete("1.0", tk.END)
        self.pdf_text_area.insert(tk.END, f"正在处理 PDF: {pdf_path}\n")
        self.pdf_analysis_results = []  # 重置分析结果

        def update_status(message):
            """更新状态信息到GUI"""
            self.pdf_text_area.insert(tk.END, message + "\n")
            self.pdf_text_area.see(tk.END)
            self.root.update()

        def process_in_thread():
            try:
                update_status("开始提取PDF页面...")
                images_data = self.extract_images_from_pdf(pdf_path)

                if not images_data:
                    update_status("[失败] PDF转换失败，未能生成图像")
                    return
                
                update_status(f"[成功] 成功转换 {len(images_data)} 页PDF为图像")
                
                for idx, image_data in enumerate(images_data, 1):
                    try:
                        update_status(f"\n[处理中] 正在分析第 {image_data['page']} 页... ({idx}/{len(images_data)})")
                        
                        analysis_result = self.analyze_image(image_data['bytes'])
                        self.pdf_analysis_results.append((image_data['page'], analysis_result, image_data['bytes']))
                        update_status(f"[成功] 第 {image_data['page']} 页分析完成")
                        update_status(f"分析结果：\n{analysis_result}\n")

                        # 显示图片
                        image_bytes_io = io.BytesIO(image_data['bytes'])
                        img = Image.open(image_bytes_io)
                        display_width = 300
                        ratio = display_width / img.width
                        display_height = int(img.height * ratio)
                        img = img.resize((display_width, display_height), Image.Resampling.LANCZOS)
                        photo = ImageTk.PhotoImage(img)

                        def add_image_to_text():
                            image_label = tk.Label(self.pdf_text_area, image=photo)
                            image_label.image = photo
                            self.pdf_text_area.window_create(tk.END, window=image_label)
                            self.pdf_text_area.insert(tk.END, "\n")
                            self.pdf_text_area.see(tk.END)

                        self.root.after(0, add_image_to_text)

                    except Exception as e:
                        update_status(f"[错误] 处理第 {image_data['page']} 页时出错：{str(e)}")
                        import traceback
                        update_status(f"详细错误：{traceback.format_exc()}")

                update_status("\n[完成] 所有页面处理完成！可以点击'下载分析结果'保存。")

            except Exception as e:
                update_status(f"[错误] 处理过程出错：{str(e)}")
                import traceback
                update_status(f"详细错误：{traceback.format_exc()}")

        threading.Thread(target=process_in_thread, daemon=True).start()

    def analyze_image(self, image_data):
        """调用大模型 API 分析图像，返回分析结果"""
        headers = {
            "Content-Type": "application/json",
        }
        
        # 将图像数据编码为 base64 字符串
        image_base64 = base64.b64encode(image_data).decode('utf-8')
        
        # Gemini 2.0 Flash API 格式
        data = {
            "contents": [
                {
                    "parts": [
                        {
                            "text": "请用中文分析这张图片的内容，精简且精准，不超过50个字"
                        },
                        {
                            "inlineData": {
                                "mimeType": "image/png",
                                "data": image_base64
                            }
                        }
                    ]
                }
            ],
            "generationConfig": {
                "temperature": 0.7,
                "maxOutputTokens": 2000,
                "topP": 1,
                "topK": 32
            },
            "safetySettings": [
                {
                    "category": "HARM_CATEGORY_HARASSMENT",
                    "threshold": "BLOCK_NONE"
                },
                {
                    "category": "HARM_CATEGORY_HATE_SPEECH",
                    "threshold": "BLOCK_NONE"
                },
                {
                    "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                    "threshold": "BLOCK_NONE"
                },
                {
                    "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
                    "threshold": "BLOCK_NONE"
                }
            ]
        }
        
        max_retries = 3
        retry_delay = 2  # 秒
        
        for attempt in range(max_retries):
            try:
                # 使用更长的超时时间
                response = requests.post(
                    f"{GOOGLE_API_URL}?key={GOOGLE_API_KEY}",
                    headers=headers,
                    json=data,
                    timeout=60,  # 增加超时时间到60秒
                    verify=True  # 确保SSL验证
                )
                
                if response.status_code != 200:
                    error_message = f"API 调用失败：HTTP {response.status_code}\n"
                    try:
                        error_detail = response.json()
                        error_message += f"错误详情：{error_detail}\n"
                    except:
                        pass
                    
                    if attempt < max_retries - 1:  # 如果还有重试机会
                        time.sleep(retry_delay)  # 等待一段时间后重试
                        continue
                    return error_message
                
                result = response.json()
                if "candidates" in result and len(result["candidates"]) > 0:
                    return result["candidates"][0]["content"]["parts"][0]["text"]
                else:
                    return "API 返回格式异常"
                
            except requests.exceptions.ConnectionError as e:
                error_message = f"连接错误 (尝试 {attempt + 1}/{max_retries}): {str(e)}"
                if attempt < max_retries - 1:
                    time.sleep(retry_delay)
                    continue
                return error_message
                
            except requests.exceptions.RequestException as e:
                return f"API 调用失败：{str(e)}"
            except Exception as e:
                return f"处理响应时出错：{str(e)}"
        
        return "所有重试都失败了"

    def extract_images_from_pdf(self, pdf_path):
        """将 PDF 的每一页转换为图片"""
        images = []
        try:
            doc = fitz.open(pdf_path)
            print(f"PDF页数: {doc.page_count}")  # 调试信息
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                # 将页面渲染为图片，使用RGB格式
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                
                # 直接从pixmap获取PNG格式的图片数据
                img_bytes = pix.tobytes("png")
                
                print(f"成功转换第 {page_num + 1} 页")  # 调试信息
                
                images.append({
                    "page": page_num + 1,
                    "index": 1,
                    "bytes": img_bytes
                })
            doc.close()
            print(f"总共转换了 {len(images)} 页")  # 调试信息
        except Exception as e:
            print(f"PDF转换失败，错误信息：{e}")
            import traceback
            print(traceback.format_exc())  # 打印详细错误信息
        return images

    def save_analysis(self):
        """保存分析结果到Word文档"""
        if not self.pdf_analysis_results:
            self.pdf_text_area.insert(tk.END, "没有可供下载的分析结果\n")
            return

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_filename = f"analysis_{timestamp}.docx"
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=default_filename,
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return

        try:
            doc = Document()
            doc.add_heading('PDF文档分析报告', 0)

            for page_num, result, image in self.pdf_analysis_results:
                doc.add_heading(f'第 {page_num} 页分析', level=1)
                doc.add_paragraph(result)
                
                # 修改图片处理逻辑
                if image:
                    try:
                        # 确保图片是PIL Image对象
                        if isinstance(image, bytes):
                            image = Image.open(io.BytesIO(image))
                        
                        # 保存为PNG格式
                        img_byte_arr = io.BytesIO()
                        image.save(img_byte_arr, format='PNG')
                        img_byte_arr.seek(0)
                        
                        # 添加图片到文档
                        doc.add_picture(img_byte_arr, width=Inches(6))
                        doc.add_paragraph('')  # 添加一个空行
                    except Exception as img_error:
                        doc.add_paragraph(f'[图片处理失败: {str(img_error)}]')
                
                doc.add_paragraph('_' * 50)

            doc.save(file_path)
            self.pdf_text_area.insert(tk.END, f"\n分析结果已保存至: {file_path}\n")
            
        except Exception as e:
            self.pdf_text_area.insert(tk.END, f"\n保存文件失败: {str(e)}\n")
            import traceback
            self.pdf_text_area.insert(tk.END, f"详细错误：{traceback.format_exc()}\n")

    def merge_pdf_md(self):
        """合并 PDF 图片和 MD 文档内容"""
        if not self.pdf_analysis_results or not self.md_analysis_results:
            messagebox.showwarning("警告", "请确保已经处理了 PDF 和 MD 文件")
            return

        # 创建预览窗口
        preview_window = tk.Toplevel(self.root)
        preview_window.title("合并预览")
        preview_window.geometry("800x600")

        # 创建预览区域
        preview_frame = tk.Frame(preview_window)
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # 创建滚动文本区域
        preview_text = scrolledtext.ScrolledText(preview_frame, wrap=tk.WORD)
        preview_text.pack(fill=tk.BOTH, expand=True)

        # 创建保存按钮
        save_button = tk.Button(preview_window, text="保存为Word文档", 
                               command=lambda: self.save_merged_document(preview_text.get("1.0", tk.END)))
        save_button.pack(pady=10)

        # 合并内容并显示
        try:
            for i, (md_page, md_result, _) in enumerate(self.md_analysis_results):
                # 添加 MD 内容
                preview_text.insert(tk.END, f"\n=== 第 {md_page} 段 ===\n\n")
                preview_text.insert(tk.END, f"{md_result}\n\n")

                # 查找对应的 PDF 页面
                for pdf_page, _, image_bytes in self.pdf_analysis_results:
                    if pdf_page == md_page:
                        # 显示图片
                        if image_bytes:
                            img = Image.open(io.BytesIO(image_bytes))
                            # 调整图片大小
                            display_width = 500
                            ratio = display_width / img.width
                            display_height = int(img.height * ratio)
                            img = img.resize((display_width, display_height), Image.Resampling.LANCZOS)
                            photo = ImageTk.PhotoImage(img)

                            # 创建图片标签
                            image_label = tk.Label(preview_text, image=photo)
                            image_label.image = photo  # 保持引用
                            preview_text.window_create(tk.END, window=image_label)
                            preview_text.insert(tk.END, "\n\n")
                        break

                preview_text.insert(tk.END, "-" * 80 + "\n\n")

        except Exception as e:
            messagebox.showerror("错误", f"合并预览失败: {str(e)}")

    def save_merged_document(self, preview_content):
        """保存合并后的文档"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"merged_document_{timestamp}.docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if not file_path:
            return

        try:
            doc = Document()
            doc.add_heading('合并文档', 0)

            for i, (md_page, md_result, _) in enumerate(self.md_analysis_results):
                # 添加 MD 内容
                doc.add_heading(f'第 {md_page} 段', level=1)
                doc.add_paragraph(md_result)

                # 添加对应的 PDF 图片
                for pdf_page, _, image_bytes in self.pdf_analysis_results:
                    if pdf_page == md_page and image_bytes:
                        img_stream = io.BytesIO(image_bytes)
                        doc.add_picture(img_stream, width=Inches(6))
                        doc.add_paragraph()  # 添加空行
                        break

                doc.add_paragraph('_' * 50)

            doc.save(file_path)
            messagebox.showinfo("成功", f"文档已保存至: {file_path}")
            
        except Exception as e:
            messagebox.showerror("错误", f"保存文档失败: {str(e)}")

    def calculate_similarity_with_api(self, text1, text2):
        """使用Google Gemini API计算文本相似度"""
        headers = {
            "Content-Type": "application/json",
        }
        
        data = {
            "contents": [{
                "parts":[{
                    "text": f"""请分析以下两段文本的相似度,返回一个0到1之间的数字(1表示完全相似,0表示完全不相似):

                    文本1:
                    {text1}

                    文本2:
                    {text2}

                    只返回相似度数值,不要有任何其他文字。"""
                }]
            }],
            "generationConfig": {
                "temperature": 0.1,  # 使用较低的temperature以获得更确定的结果
                "maxOutputTokens": 100,
            }
        }
        
        try:
            response = requests.post(
                f"{GOOGLE_API_URL}?key={GOOGLE_API_KEY}",
                headers=headers,
                json=data,
                timeout=30
            )
            response.raise_for_status()
            
            if response.status_code == 200:
                result = response.json()
                if "candidates" in result and len(result["candidates"]) > 0:
                    similarity = float(result["candidates"][0]["content"]["parts"][0]["text"].strip())
                    return similarity
        except Exception as e:
            print(f"计算相似度时出错: {str(e)}")
        return 0.0  # 出错时返回0

    def intelligently_match_content(self, similarity_threshold=0.6):
        """智能匹配PDF图片和MD段落"""
        if not self.pdf_analysis_results or not self.md_analysis_results:
            messagebox.showwarning("警告", "请确保已经处理了 PDF 和 MD 文件")
            return

        self.pdf_text_area.insert(tk.END, "正在进行智能内容匹配...\n")
        self.matched_content = []

        # 为每个PDF页面找到最匹配的MD段落
        for pdf_idx, (pdf_page, pdf_desc, pdf_image) in enumerate(self.pdf_analysis_results):
            best_match_score = 0
            best_match_md = None
            
            # 计算与每个MD段落的相似度
            for md_idx, (md_page, md_content, _) in enumerate(self.md_analysis_results):
                similarity = self.calculate_similarity_with_api(pdf_desc, md_content)
                
                if similarity > best_match_score:
                    best_match_score = similarity
                    best_match_md = (md_page, md_content)
            
            # 记录匹配结果
            match_info = {
                "pdf_page": pdf_page,
                "pdf_description": pdf_desc,
                "pdf_image_bytes": pdf_image,
                "md_page": best_match_md[0] if best_match_score >= similarity_threshold else None,
                "md_content": best_match_md[1] if best_match_score >= similarity_threshold else None,
                "similarity_score": best_match_score
            }
            
            self.matched_content.append(match_info)
            
            # 显示匹配结果
            if best_match_score >= similarity_threshold:
                self.pdf_text_area.insert(tk.END, 
                    f"PDF第{pdf_page}页 匹配到 MD第{best_match_md[0]}段 (相似度: {best_match_score:.2f})\n")
            else:
                self.pdf_text_area.insert(tk.END, 
                    f"PDF第{pdf_page}页 未找到足够相似的MD段落\n")
        
        self.pdf_text_area.insert(tk.END, "智能内容匹配完成。\n")
        messagebox.showinfo("完成", "智能内容匹配完成！")

    def __init__(self, root):
        self.root = root
        self.root.title("文档分析工具")
        self.pdf_analysis_results = []
        self.md_analysis_results = []

        # 创建单个按钮框架
        button_frame = tk.Frame(self.root)
        button_frame.pack(pady=5)

        # 所有按钮放在同一行
        self.btn_select_pdf = tk.Button(button_frame, text="选择 PDF 文件", command=self.process_pdf)
        self.btn_select_pdf.pack(side=tk.LEFT, padx=5)

        self.btn_download_pdf = tk.Button(button_frame, text="下载PDF分析", command=self.save_analysis)
        self.btn_download_pdf.pack(side=tk.LEFT, padx=5)

        self.btn_select_md = tk.Button(button_frame, text="选择 MD 文件", command=self.process_md)
        self.btn_select_md.pack(side=tk.LEFT, padx=5)

        self.btn_download_md = tk.Button(button_frame, text="下载MD分析", command=self.save_md_analysis)
        self.btn_download_md.pack(side=tk.LEFT, padx=5)

        # 添加合并按钮
        self.btn_merge = tk.Button(button_frame, text="合并预览", command=self.merge_pdf_md)
        self.btn_merge.pack(side=tk.LEFT, padx=5)

        # 创建分隔的文本区域
        self.paned_window = tk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        self.paned_window.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

        # PDF 文本区域
        self.pdf_text_area = scrolledtext.ScrolledText(self.paned_window, height=20, wrap=tk.WORD)
        self.paned_window.add(self.pdf_text_area)

        # Markdown 文本区域
        self.md_text_area = scrolledtext.ScrolledText(self.paned_window, height=20, wrap=tk.WORD)
        self.paned_window.add(self.md_text_area)

if __name__ == "__main__":
    # API配置
    GOOGLE_API_KEY = "你的密钥"  # Google AI API 密钥
    GOOGLE_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash-exp:generateContent"
    
    root = tk.Tk()
    app = PDFAnalyzer(root)
    root.mainloop()

