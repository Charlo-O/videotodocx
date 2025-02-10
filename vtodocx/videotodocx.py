import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox
import requests
from PIL import Image
import io
from datetime import datetime
from docx import Document
import threading
import markdown
from docx.shared import Inches
import base64
from PIL import Image, ImageTk  # 添加 ImageTk
import time as time_module  # 重命名time模块避免冲突
from moviepy import VideoFileClip, ImageClip
import cv2
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
import os  # 导入 os 模块

# API配置 (请替换为你自己的密钥)
GOOGLE_API_KEY = "请替换为你自己的密钥"
GOOGLE_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash-lite-preview-02-05:generateContent"

# 预训练的 Sentence Transformer 模型 (用于计算文本相似度)
MODEL_NAME = "all-mpnet-base-v2"  # 这是一个常用的、效果较好的模型
SIMILARITY_THRESHOLD = 0.85  # 相似度阈值，用于判断帧是否相似

class APILimiter:
    def __init__(self, requests_per_minute=60):  # 默认每分钟60个请求
        self.requests_per_minute = requests_per_minute
        self.requests = []
        self.lock = threading.Lock()

    def wait_if_needed(self):
        current_time = time_module.time()
        with self.lock:
            # 移除超过1分钟的请求记录
            self.requests = [req_time for req_time in self.requests
                           if current_time - req_time < 60]

            if len(self.requests) >= self.requests_per_minute:
                # 计算需要等待的时间
                wait_time = 60 - (current_time - self.requests[0])
                if wait_time > 0:
                    time_module.sleep(wait_time)
                self.requests = self.requests[1:]  # 移除最老的请求

            self.requests.append(current_time)

class VideoAnalyzer:
    def __init__(self, root):
        self.root = root
        self.root.title("视频文本分析工具 (以文搜图)")
        self.video_path = None  # 存储视频路径
        self.video_clip = None # 视频剪辑
        self.md_analysis_results = []  # 存储 Markdown 分析结果
        self.sentence_model = SentenceTransformer(MODEL_NAME)  # 加载 Sentence Transformer 模型
        self.frame_embeddings = {}  # 存储帧嵌入
        self.unique_frame_embeddings = {} # 存储去重后的帧嵌入和时间
        self.unique_frames_data = [] # 存储去重后的帧数据 (PIL Image对象)
        self.api_limiter = APILimiter()  # 添加 API 限制器实例
        self.md_content = "" # 存储markdown 文件内容


    def process_md(self):
        """选择 Markdown 文档，读取内容，但不立即处理分析"""
        md_path = filedialog.askopenfilename(filetypes=[("Markdown Files", "*.md")])
        if not md_path:
            return

        self.md_text_area.delete("1.0", tk.END)
        self.md_text_area.insert(tk.END, f"已选择 Markdown: {md_path}\n")

        try:
            with open(md_path, 'r', encoding='utf-8') as file:
                self.md_content = file.read()
        except Exception as e:
             self.md_text_area.insert(tk.END, f"读取 Markdown 文件失败: {str(e)}\n")
             self.md_content = ""


    def process_video(self):
        """选择视频文件 (仅选择，不立即处理)"""
        video_path = filedialog.askopenfilename(filetypes=[("Video Files", "*.mp4")])
        if not video_path:
            return

        self.video_path = video_path
        self.video_text_area.delete("1.0", tk.END)
        self.video_text_area.insert(tk.END, f"已选择视频: {video_path}\n")
        # 加载视频 (如果尚未加载)
        if self.video_clip is None and self.video_path:
            self.video_clip = VideoFileClip(self.video_path)

    def generate_unique_frame_embeddings(self):
        """生成视频帧的嵌入向量并去重"""
        if not self.video_path:
            messagebox.showwarning("警告", "请先选择视频文件")
            return
        if not self.video_clip:
            messagebox.showwarning("警告", "视频加载失败, 重新选择")
            return

        self.video_text_area.delete("1.0", tk.END)  #清空
        self.video_text_area.insert(tk.END, "正在生成视频帧嵌入并去重...\n")
        self.frame_embeddings = {}  # 清空之前的嵌入
        self.unique_frame_embeddings = {} # 清空
        self.unique_frames_data = [] # 清空

        def generate_embeddings_thread():
            try:
                duration = self.video_clip.duration
                fps = self.video_clip.fps
                frame_interval = 1.5  # 每 1.5 秒采样一帧 (可调整)
                total_frames = int(duration / frame_interval)
                processed_frames = 0
                unique_frame_count = 0
                unique_embeddings_list = [] # 存储唯一帧的embeddings

                output_folder = "unique_frames"  # 保存唯一帧的文件夹
                if not os.path.exists(output_folder):
                    os.makedirs(output_folder)

                for t in np.arange(0, duration, frame_interval):
                    frame = self.video_clip.to_ImageClip(t).img  # 使用 to_ImageClip 获取 PIL Image
                    img = Image.fromarray(frame)
                    img_bytes = io.BytesIO()
                    img.save(img_bytes, format='PNG')
                    img_bytes_value = img_bytes.getvalue()

                    # 使用 Gemini API 分析图像并获取文本描述 (更详细的提示)
                    image_desc = self.analyze_image(img_bytes_value, time=t)

                    # 使用 Sentence Transformer 计算文本描述的嵌入
                    embedding = self.sentence_model.encode(image_desc, convert_to_tensor=True)


                    is_unique = True
                    if unique_embeddings_list:
                        for unique_embedding in unique_embeddings_list:
                            similarity = cosine_similarity(embedding.cpu().reshape(1, -1), unique_embedding.cpu().reshape(1, -1))[0][0]
                            if similarity >= SIMILARITY_THRESHOLD:
                                is_unique = False
                                break # 找到相似帧，跳出循环

                    if is_unique:
                        unique_embeddings_list.append(embedding)
                        self.unique_frame_embeddings[t] = embedding # 存储唯一帧的embedding 和 时间
                        self.unique_frames_data.append(img) # 存储PIL Image 对象
                        unique_frame_count += 1

                        # 保存唯一帧到文件
                        frame_filename = os.path.join(output_folder, f"frame_{unique_frame_count}_{t:.2f}s.png")
                        img.save(frame_filename)
                        self.video_text_area.insert(tk.END, f"保存唯一帧: {frame_filename}\n")


                    self.frame_embeddings[t] = embedding  # 仍然保存所有帧的embedding，方便后续使用 (如果需要)

                    processed_frames += 1
                    progress = (processed_frames / total_frames) * 100
                    self.video_text_area.insert(tk.END, f"已处理 {processed_frames}/{total_frames} 帧 ({progress:.2f}%). 唯一帧数量: {unique_frame_count}\n")
                    self.video_text_area.see(tk.END)  # 滚动到末尾
                    self.root.update()

                self.video_text_area.insert(tk.END, f"帧嵌入生成和去重完成！共提取 {unique_frame_count} 个唯一帧.\n")
                self.video_text_area.insert(tk.END, f"唯一帧保存在 '{output_folder}' 文件夹中\n")


            except Exception as e:
                self.video_text_area.insert(tk.END, f"生成嵌入时出错: {str(e)}\n")
                import traceback
                self.video_text_area.insert(tk.END, f"详细错误：{traceback.format_exc()}\n")

        threading.Thread(target=generate_embeddings_thread, daemon=True).start()


    def segment_markdown_text(self):
        """使用Gemini API智能分析和分割Markdown文本"""
        if not self.md_content:
            messagebox.showwarning("警告", "请先选择 Markdown 文件")
            return
        if not self.unique_frames_data:
            messagebox.showwarning("警告", "请先生成并去重视频帧")
            return

        self.md_text_area.delete("1.0", tk.END)
        num_unique_frames = len(self.unique_frames_data)
        self.md_text_area.insert(tk.END, f"正在使用AI分析并分割文本...\n")
        self.md_analysis_results = [] # 清空之前的分析结果

        def segment_thread():
            try:
                # 构建API请求
                prompt = f"""请对以下文本进行优化和整理，使其更加清晰和准确。联系上下文，纠正里面的错别字，符合事实，前后文相互呼应，不要随便删除内容，合理运用标点符号，使其更加通顺。
                根据文本内容，将其分成{num_unique_frames}个段落，每个段落的内容要完整且相关。
                直接返回优化后的文本内容，不要添加任何额外的解释或标记。之后按照如下格式输出结果：

                段落 1:
                [文本内容]
                关键词: [关键词列表]
                关键实体: [实体列表]
                时间信息: [时间相关信息]
                场景描述: [场景细节描述]

                段落 2:
                [依此类推...]

                要分析的文本：
                {self.md_content}
                """

                # 调用Gemini API
                headers = {
                    "Content-Type": "application/json",
                    "x-goog-api-key": GOOGLE_API_KEY
                }
                
                data = {
                    "contents": [{
                        "parts": [{
                            "text": prompt
                        }]
                    }]
                }

                self.api_limiter.wait_if_needed()  # 使用API限制器
                response = requests.post(GOOGLE_API_URL, headers=headers, json=data)
                
                if response.status_code == 200:
                    result = response.json()
                    if 'candidates' in result and result['candidates']:
                        analysis_text = result['candidates'][0]['content']['parts'][0]['text']
                        
                        # 解析API返回的结果
                        segments = []
                        current_segment = {}
                        lines = analysis_text.strip().split('\n')
                        
                        for line in lines:
                            line = line.strip()
                            if line.startswith('段落'):
                                if current_segment:
                                    segments.append(current_segment)
                                current_segment = {'segment_num': len(segments) + 1}
                            elif line.startswith('关键词:'):
                                current_segment['keywords'] = line[4:].strip()
                            elif line.startswith('关键实体:'):
                                current_segment['entities'] = line[5:].strip()
                            elif line.startswith('时间信息:'):
                                current_segment['time_info'] = line[5:].strip()
                            elif line.startswith('场景描述:'):
                                current_segment['scene_desc'] = line[5:].strip()
                            elif line and current_segment:
                                if 'text' not in current_segment:
                                    current_segment['text'] = line
                                else:
                                    current_segment['text'] += '\n' + line
                        
                        if current_segment:
                            segments.append(current_segment)
                        
                        # 保存分析结果
                        self.md_analysis_results = []
                        for segment in segments:
                            segment_info = {
                                "segment_num": segment.get('segment_num', 0),
                                "text": segment.get('text', '').strip(),
                                "keywords": segment.get('keywords', ''),
                                "entities": segment.get('entities', ''),
                                "time_info": segment.get('time_info', ''),
                                "scene_desc": segment.get('scene_desc', ''),
                                "frame_data": None
                            }
                            self.md_analysis_results.append(segment_info)
                            
                            # 显示分析结果
                            self.md_text_area.insert(tk.END, f"\n=== 分割段落 {segment_info['segment_num']} ===\n")
                            self.md_text_area.insert(tk.END, f"内容: {segment_info['text'][:200]}...\n")
                            self.md_text_area.insert(tk.END, f"关键词: {segment_info['keywords']}\n")
                            self.md_text_area.insert(tk.END, f"实体: {segment_info['entities']}\n")
                            self.md_text_area.insert(tk.END, f"时间: {segment_info['time_info']}\n")
                            self.md_text_area.insert(tk.END, f"场景: {segment_info['scene_desc']}\n")
                        
                        self.md_text_area.insert(tk.END, f"\nAI分析完成！已将文本分割成 {len(segments)} 段.\n")
                    else:
                        self.md_text_area.insert(tk.END, "API返回结果格式错误\n")
                else:
                    self.md_text_area.insert(tk.END, f"API调用失败: {response.status_code}\n{response.text}\n")

            except Exception as e:
                self.md_text_area.insert(tk.END, f"分析文本失败: {str(e)}\n")
                import traceback
                self.md_text_area.insert(tk.END, f"详细错误：{traceback.format_exc()}\n")

        threading.Thread(target=segment_thread, daemon=True).start()


    def match_text_to_frames(self):
        """根据分割后的 Markdown 文本，匹配视频帧 (这里修改为匹配唯一帧)"""
        if not self.md_analysis_results:
            messagebox.showwarning("警告", "请先分割 Markdown 文本")
            return
        if not self.unique_frame_embeddings:
            messagebox.showwarning("警告", "请先生成并去重视频帧")
            return
        if not self.unique_frames_data:
            messagebox.showwarning("警告", "唯一帧数据丢失，请重新生成")
            return

        if len(self.md_analysis_results) != len(self.unique_frames_data):
            messagebox.showwarning("警告", f"文本段落数 ({len(self.md_analysis_results)}) 与唯一帧数 ({len(self.unique_frames_data)}) 不匹配。请检查处理流程。")
            return


        self.video_text_area.delete("1.0", tk.END) # 清空
        self.video_text_area.insert(tk.END, "正在将文本段落匹配到唯一视频帧...\n")

        def match_in_thread():
            try:
                for i, segment_info in enumerate(self.md_analysis_results):
                    if i >= len(self.unique_frames_data): # 防止索引越界，虽然理论上不应该发生
                        self.video_text_area.insert(tk.END, f"\n警告：段落 {segment_info['segment_num']} 没有对应的唯一帧 (帧数不足).\n")
                        continue

                    # 使用分割后的段落文本进行匹配
                    match_text = segment_info['text']

                    # 计算匹配文本的嵌入
                    text_embedding = self.sentence_model.encode(match_text, convert_to_tensor=True)

                    #  使用预先计算好的唯一帧嵌入
                    unique_frame_times = list(self.unique_frame_embeddings.keys())
                    best_match_time = unique_frame_times[i] # 假设段落顺序对应唯一帧顺序
                    best_frame_embedding = self.unique_frame_embeddings[best_match_time]
                    similarity = cosine_similarity(text_embedding.cpu().reshape(1, -1), best_frame_embedding.cpu().reshape(1, -1))[0][0]


                    # 获取对应的唯一帧图像数据
                    best_frame_image = self.unique_frames_data[i] #  假设顺序一致
                    img_bytes = io.BytesIO()
                    best_frame_image.save(img_bytes, format='PNG')
                    segment_info["frame_data"] = img_bytes.getvalue()  # 存储帧数据


                    # 显示匹配结果 (在界面上)
                    self.video_text_area.insert(tk.END, f"\n段落 {segment_info['segment_num']} 匹配到 唯一帧: {best_match_time:.2f} 秒\n")
                    self.video_text_area.insert(tk.END, f"相似度: {similarity:.4f}\n")

                    # 显示图像
                    img = Image.open(io.BytesIO(segment_info["frame_data"]))
                    display_width = 300
                    ratio = display_width / img.width
                    display_height = int(img.height * ratio)
                    img = img.resize((display_width, display_height), Image.Resampling.LANCZOS)
                    photo = ImageTk.PhotoImage(img)
                    image_label = tk.Label(self.video_text_area, image=photo)
                    image_label.image = photo  # 保持引用
                    self.video_text_area.window_create(tk.END, window=image_label)
                    self.video_text_area.insert(tk.END, "\n")
                    self.video_text_area.see(tk.END)


                self.video_text_area.insert(tk.END, "文本段落与唯一帧匹配完成！\n")

            except Exception as e:
                self.video_text_area.insert(tk.END, f"匹配过程中出错: {str(e)}\n")
                import traceback
                self.video_text_area.insert(tk.END, f"详细错误：{traceback.format_exc()}\n")

        threading.Thread(target=match_in_thread, daemon=True).start()



    def analyze_image(self, image_data, time=None):
        """调用 Gemini API 分析图像 (更具体的提示)"""
        headers = {"Content-Type": "application/json"}
        image_base64 = base64.b64encode(image_data).decode('utf-8')

        prompt = f"详细描述这张图片的内容，重点关注时间 {time:.2f} 秒 附近的事件、物体、人物及其动作和环境。" if time else "详细描述这张图片的内容。"
        prompt += " 尽可能提供细节，例如颜色、形状、纹理、相对位置等。"

        data = {
            "contents": [{
                "parts": [
                    {"text": prompt},
                    {
                        "inlineData": {
                            "mimeType": "image/png",
                            "data": image_base64
                        }
                    }
                ]
            }],
            "generationConfig": {
                "temperature": 0.4,
                "maxOutputTokens": 500,  # 增加输出长度
                "topP": 1,
                "topK": 32
            },
            "safetySettings": [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
            ]
        }

        max_retries = 3
        retry_delay = 2

        for attempt in range(max_retries):
            try:
                response = requests.post(
                    f"{GOOGLE_API_URL}?key={GOOGLE_API_KEY}",
                    headers=headers,
                    json=data,
                    timeout=60,
                    verify=True
                )

                if response.status_code != 200:
                    error_message = f"API 调用失败：HTTP {response.status_code}\n"
                    try:
                        error_detail = response.json()
                        error_message += f"错误详情：{error_detail}\n"
                    except:
                        pass
                    if attempt < max_retries - 1:
                        time_module.sleep(retry_delay)  # 使用重命名后的time模块
                        continue
                    return error_message

                result = response.json()
                if "candidates" in result and len(result["candidates"]) > 0:
                    return result["candidates"][0]["content"]["parts"][0]["text"]
                else:
                    return "API 返回格式异常"

            except requests.exceptions.ConnectionError as e:
                error_message = f"连接错误 (尝试 {attempt + 1}/{max_retries}): {str(e)}"
                if attempt < max_retries - 1:
                    time_module.sleep(retry_delay)  # 使用重命名后的time模块
                    continue
                return error_message

            except requests.exceptions.RequestException as e:
                return f"API 调用失败：{str(e)}"
            except Exception as e:
                return f"处理响应时出错：{str(e)}"

        return "所有重试都失败了"

    def save_merged_document(self):
        """保存合并后的文档 (Markdown 分析结果 + 匹配的帧)"""
        if not self.md_analysis_results:
            messagebox.showwarning("警告", "没有可保存的内容。请先处理 Markdown 文件并匹配视频帧。")
            return

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"merged_document_{timestamp}.docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )

        if not file_path:
            return

        try:
            doc = Document()
            doc.add_heading('Markdown 与视频帧匹配结果', 0)

            for segment_info in self.md_analysis_results:
                doc.add_heading(f"段落 {segment_info['segment_num']}", level=1)
                doc.add_paragraph(f"关键词: {segment_info['keywords']}")
                doc.add_paragraph(f"关键实体: {segment_info['entities']}")
                doc.add_paragraph(f"时间信息: {segment_info['time_info']}")
                doc.add_paragraph(f"场景描述: {segment_info['scene_desc']}")
                doc.add_paragraph(segment_info["text"])

                if segment_info["frame_data"]:
                    img_stream = io.BytesIO(segment_info["frame_data"])
                    doc.add_picture(img_stream, width=Inches(6))
                else:
                    doc.add_paragraph("(未找到匹配的帧)")

                doc.add_paragraph("-" * 50)

            doc.save(file_path)
            messagebox.showinfo("成功", f"文档已保存至: {file_path}")

        except Exception as e:
            messagebox.showerror("错误", f"保存文档失败: {str(e)}")


    def show_preview(self):
        """显示合并预览 (Markdown 分析结果 + 匹配的帧)"""
        if not self.md_analysis_results:
            messagebox.showwarning("警告", "没有可预览的内容。请先处理 Markdown 文件并匹配视频帧。")
            return

        preview_window = tk.Toplevel(self.root)
        preview_window.title("合并预览")
        preview_window.geometry("800x600")

        preview_frame = tk.Frame(preview_window)
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        preview_text = scrolledtext.ScrolledText(preview_frame, wrap=tk.WORD)
        preview_text.pack(fill=tk.BOTH, expand=True)

        # 添加保存按钮
        save_button = tk.Button(preview_window, text="保存为Word文档", command=self.save_merged_document)
        save_button.pack(pady=10)

        try:
            for segment_info in self.md_analysis_results:
                preview_text.insert(tk.END, f"\n=== 段落 {segment_info['segment_num']} ===\n")
                preview_text.insert(tk.END, f"关键词: {segment_info['keywords']}\n")
                preview_text.insert(tk.END, f"关键实体: {segment_info['entities']}\n")
                preview_text.insert(tk.END, f"时间信息: {segment_info['time_info']}\n")
                preview_text.insert(tk.END, f"场景描述: {segment_info['scene_desc']}\n")
                preview_text.insert(tk.END, f"{segment_info['text']}\n\n")

                if segment_info["frame_data"]:
                    img = Image.open(io.BytesIO(segment_info["frame_data"]))
                    display_width = 500
                    ratio = display_width / img.width
                    display_height = int(img.height * ratio)
                    img = img.resize((display_width, display_height), Image.Resampling.LANCZOS)
                    photo = ImageTk.PhotoImage(img)
                    image_label = tk.Label(preview_text, image=photo)
                    image_label.image = photo  # 保持引用
                    preview_text.window_create(tk.END, window=image_label)
                    preview_text.insert(tk.END, "\n")
                else:
                    preview_text.insert(tk.END, "(未找到匹配的帧)\n")

                preview_text.insert(tk.END, "-" * 80 + "\n")

        except Exception as e:
            messagebox.showerror("错误", f"预览生成失败: {str(e)}")



# ... (其余部分与之前的代码结构类似，只需调整按钮和布局)
    def create_widgets(self):
        """创建 GUI 控件"""
        button_frame = tk.Frame(self.root)
        button_frame.pack(pady=5)

        # 修改按钮顺序和文本
        self.btn_select_video = tk.Button(button_frame, text="1. 选择视频文件", command=self.process_video)
        self.btn_select_video.pack(side=tk.LEFT, padx=5)

        self.btn_gen_embeddings = tk.Button(button_frame, text="2. 生成 & 去重帧", command=self.generate_unique_frame_embeddings)
        self.btn_gen_embeddings.pack(side=tk.LEFT, padx=5)

        self.btn_select_md = tk.Button(button_frame, text="3. 选择 MD 文件", command=self.process_md)
        self.btn_select_md.pack(side=tk.LEFT, padx=5)

        self.btn_segment_md = tk.Button(button_frame, text="4. 分割 MD 文本", command=self.segment_markdown_text)
        self.btn_segment_md.pack(side=tk.LEFT, padx=5)

        self.btn_match = tk.Button(button_frame, text="5. 匹配文本与帧", command=self.match_text_to_frames)
        self.btn_match.pack(side=tk.LEFT, padx=5)

        self.btn_preview = tk.Button(button_frame, text="6. 预览", command=self.show_preview)  # 新增预览按钮
        self.btn_preview.pack(side=tk.LEFT, padx=5)


        self.paned_window = tk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        self.paned_window.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

        self.md_text_area = scrolledtext.ScrolledText(self.paned_window, height=20, wrap=tk.WORD)
        self.paned_window.add(self.md_text_area)

        self.video_text_area = scrolledtext.ScrolledText(self.paned_window, height=20, wrap=tk.WORD)
        self.paned_window.add(self.video_text_area)


    def __init__(self, root):
        self.root = root
        self.root.title("视频文本分析工具 (以文搜图)")
        self.video_path = None
        self.video_clip = None
        self.md_analysis_results = []
        self.sentence_model = SentenceTransformer(MODEL_NAME)  # Sentence Transformer 模型
        self.frame_embeddings = {}
        self.unique_frame_embeddings = {}
        self.unique_frames_data = []
        self.api_limiter = APILimiter()  # 添加 API 限制器实例
        self.md_content = ""

        self.create_widgets()

if __name__ == "__main__":
    root = tk.Tk()
    app = VideoAnalyzer(root)
    root.mainloop()
